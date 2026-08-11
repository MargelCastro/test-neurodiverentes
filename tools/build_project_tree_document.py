from pathlib import Path
from datetime import date
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "entregables"
ASSET_DIR = OUT_DIR / "_doc_assets"
OUT_FILE = OUT_DIR / "Arbol_de_procesos_y_estructura_testneurodivergentes.docx"

NAVY = "13233A"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
SLATE = "475569"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GREEN = "1F7A5A"
AMBER = "B7791F"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111827"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def font_path(bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    return next((p for p in candidates if p.exists()), candidates[0])


def load_font(size, bold=False):
    return ImageFont.truetype(str(font_path(bold)), size=size)


def wrap(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(draw, xy, title, subtitle, fill, outline, title_font, body_font):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=22, fill=fill, outline=outline, width=3)
    title_lines = wrap(draw, title, title_font, x2 - x1 - 36)
    subtitle_lines = wrap(draw, subtitle, body_font, x2 - x1 - 36)
    y = y1 + 18
    for line in title_lines:
        draw.text((x1 + 18, y), line, fill="#13233A", font=title_font)
        y += 30
    y += 6
    for line in subtitle_lines:
        draw.text((x1 + 18, y), line, fill="#475569", font=body_font)
        y += 24


def arrow(draw, start, end, color="#64748B", width=5):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) > abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 16 * sign, y2 - 9), (x2 - 16 * sign, y2 + 9)]
    else:
        sign = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - 16 * sign), (x2 + 9, y2 - 16 * sign)]
    draw.polygon(pts, fill=color)


def create_architecture_diagram(path):
    img = Image.new("RGB", (1600, 1050), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(29, True)
    body_font = load_font(23)
    heading_font = load_font(38, True)
    draw.text((60, 35), "Arquitectura funcional del sitio", fill="#13233A", font=heading_font)

    layers = [
        ("1. Experiencia publica", "Portada, directorio, guias, paginas de confianza y redirecciones", "#E8EEF5", "#2E74B5"),
        ("2. Acceso y enrutamiento", "Formulario previo: edad, audiencia, tutor y destino seguro", "#F7F9FC", "#64748B"),
        ("3. Motor de cuestionarios", "Adultos: 100 preguntas | Infantil: 60 preguntas | escala 0-4", "#FFF7E6", "#B7791F"),
        ("4. Servicios compartidos del navegador", "Navegacion, accesibilidad, foco, dialogs y almacenamiento local", "#EDF7F2", "#1F7A5A"),
        ("5. Publicacion y control", "docs/, CNAME, robots.txt, sitemap.xml, ads.txt, AdSense y pruebas", "#FDECEC", "#9B1C1C"),
    ]
    y = 115
    for idx, (title, subtitle, fill, outline) in enumerate(layers):
        rounded_box(draw, (110, y, 1490, y + 140), title, subtitle, fill, outline, title_font, body_font)
        if idx < len(layers) - 1:
            arrow(draw, (800, y + 140), (800, y + 175), color="#64748B")
        y += 175
    img.save(path, quality=95)


def create_process_diagram(path):
    img = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(26, True)
    body_font = load_font(20)
    heading_font = load_font(38, True)
    draw.text((55, 32), "Flujo principal de una evaluacion", fill="#13233A", font=heading_font)

    boxes = {
        "entry": (60, 135, 360, 285, "Entrada", "Portada, directorio o guia"),
        "select": (430, 135, 760, 285, "Seleccion", "Test adulto o infantil"),
        "form": (830, 135, 1190, 285, "Formulario previo", "Valida edad y consentimiento"),
        "quiz": (1260, 135, 1720, 285, "Cuestionario", "Carga motor y progreso guardado"),
        "answer": (1260, 430, 1720, 600, "Ciclo por pregunta", "Responder -> guardar -> avanzar / volver / pausar"),
        "complete": (830, 430, 1190, 600, "Completitud", "Todas las respuestas deben existir"),
        "score": (430, 430, 760, 600, "Calculo", "Normaliza puntaje y areas a 0-100"),
        "result": (60, 430, 360, 600, "Resultado", "Banda, perfil y resumen por areas"),
        "actions": (430, 760, 1190, 950, "Acciones posteriores", "Revisar respuestas | repetir | volver al directorio | consultar profesional"),
    }
    fills = ["#E8EEF5", "#E8EEF5", "#F7F9FC", "#FFF7E6", "#FFF7E6", "#F7F9FC", "#EDF7F2", "#EDF7F2", "#FDECEC"]
    outlines = ["#2E74B5", "#2E74B5", "#64748B", "#B7791F", "#B7791F", "#64748B", "#1F7A5A", "#1F7A5A", "#9B1C1C"]
    for (key, box), fill, outline in zip(boxes.items(), fills, outlines):
        x1, y1, x2, y2, title, subtitle = box
        rounded_box(draw, (x1, y1, x2, y2), title, subtitle, fill, outline, title_font, body_font)

    arrow(draw, (360, 210), (430, 210))
    arrow(draw, (760, 210), (830, 210))
    arrow(draw, (1190, 210), (1260, 210))
    arrow(draw, (1490, 285), (1490, 430))
    arrow(draw, (1260, 515), (1190, 515))
    arrow(draw, (830, 515), (760, 515))
    arrow(draw, (430, 515), (360, 515))
    arrow(draw, (210, 600), (520, 760))
    arrow(draw, (900, 600), (900, 760))
    img.save(path, quality=95)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Tree" not in [s.name for s in styles]:
        tree = styles.add_style("Tree", WD_STYLE_TYPE.PARAGRAPH)
        tree.base_style = styles["Normal"]
        tree.font.name = "Consolas"
        tree._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        tree._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        tree.font.size = Pt(8.7)
        tree.font.color.rgb = rgb(NAVY)
        tree.paragraph_format.space_after = Pt(0)
        tree.paragraph_format.line_spacing = 1.0

    if "Lead" not in [s.name for s in styles]:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
        lead.base_style = styles["Normal"]
        lead.font.size = Pt(12)
        lead.font.color.rgb = rgb(SLATE)
        lead.paragraph_format.space_after = Pt(10)
        lead.paragraph_format.line_spacing = 1.25

    if "Table Citation" not in [s.name for s in styles]:
        tc = styles.add_style("Table Citation", WD_STYLE_TYPE.PARAGRAPH)
        tc.base_style = styles["Normal"]
        tc.font.size = Pt(9)
        tc.font.italic = True
        tc.font.color.rgb = rgb(MUTED)
        tc.paragraph_format.space_before = Pt(4)
        tc.paragraph_format.space_after = Pt(4)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "TESTS NEURODIVERGENTES ONLINE  |  DOCUMENTACION TECNICA"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.runs[0]
    hr.font.name = "Calibri"
    hr.font.size = Pt(8.5)
    hr.font.bold = True
    hr.font.color.rgb = rgb(MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("DOCUMENTACION DEL PROYECTO")
    kr.bold = True
    kr.font.size = Pt(11)
    kr.font.color.rgb = rgb(AMBER)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Árbol de procesos\ny estructura del sitio")
    tr.bold = True
    tr.font.size = Pt(29)
    tr.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("testneurodivergentes.com")
    sr.font.size = Pt(16)
    sr.font.color.rgb = rgb(BLUE_DARK)
    subtitle.paragraph_format.space_after = Pt(50)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("Mapa tecnico, flujo de usuario, cuestionarios, datos, publicacion y control de calidad")
    lr.font.size = Pt(11)
    lr.font.italic = True
    lr.font.color.rgb = rgb(SLATE)
    line.paragraph_format.space_after = Pt(78)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Estado auditado: 10 de agosto de 2026\nRama: main | Revision observada: 29ee9d8\nAlcance: codigo local del repositorio y pruebas automatizadas")
    mr.font.size = Pt(10)
    mr.font.color.rgb = rgb(MUTED)
    doc.add_page_break()


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = rgb(accent)
    r2 = p.add_run(text)
    r2.font.color.rgb = rgb(BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, items, style="List Bullet"):
    """Create a real Word bullet list with stable preset geometry."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_numbered_list(doc, items):
    """Create a real Word decimal list that restarts at 1 for each process."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, 120)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = rgb(NAVY)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(9.3)
    set_table_geometry(table, widths_dxa, 120)
    return table


def add_tree(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="Tree")
        p.add_run(line)


def build_document():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)
    architecture_img = ASSET_DIR / "arquitectura.png"
    process_img = ASSET_DIR / "flujo_evaluacion.png"
    create_architecture_diagram(architecture_img)
    create_process_diagram(process_img)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Contenido", level=1)
    toc_items = [
        "Resumen ejecutivo y alcance",
        "Arquitectura funcional",
        "Árbol estructural del repositorio",
        "Árbol de navegación y contenidos",
        "Flujo completo de una evaluacion",
        "Formulario previo y enrutamiento",
        "Motores de cuestionarios",
        "Calculo y presentacion de resultados",
        "Persistencia, privacidad y datos",
        "Navegacion, accesibilidad y responsive",
        "SEO, publicacion y monetizacion",
        "Pruebas y control de calidad",
        "Mantenimiento y puntos de atencion",
        "Anexo A. Inventario funcional de archivos",
    ]
    add_numbered_list(doc, toc_items)
    add_callout(doc, "Lectura recomendada", "Las secciones 2 y 5 explican el funcionamiento general. Las secciones 3 y el Anexo A sirven como referencia tecnica archivo por archivo.")

    doc.add_heading("1. Resumen ejecutivo y alcance", level=1)
    doc.add_paragraph(
        "El proyecto es un sitio web estatico en espanol dedicado a informacion y cribado orientativo sobre TDA/TDAH. "
        "No utiliza un servidor de aplicaciones ni una base de datos propia: las paginas, estilos y scripts se publican directamente desde la carpeta docs/. "
        "La interactividad ocurre en el navegador del visitante."
    )
    add_callout(doc, "Resultado de la auditoria", "39 archivos publicados: 15 HTML, 9 JavaScript, 3 CSS, 8 imagenes WebP y 4 archivos de configuracion/publicacion. La suite actual ejecuta 36 comprobaciones y todas pasan.", fill="EDF7F2", accent=GREEN)
    doc.add_heading("Que hace el sitio", level=2)
    add_bullets(doc, [
        "Ofrece una portada informativa y un directorio de evaluaciones disponibles.",
        "Publica una guia central sobre TDAH y tres guias practicas relacionadas.",
        "Solicita datos minimos de acceso para dirigir al cuestionario adulto o infantil segun la audiencia y la edad.",
        "Ejecuta dos cuestionarios: 100 preguntas para adultos y 60 para poblacion infantil, con cinco opciones de frecuencia de 0 a 4.",
        "Guarda el progreso localmente en el navegador para permitir pausa, salida y recuperacion.",
        "Calcula puntajes normalizados de 0 a 100, bandas orientativas y resumenes por areas.",
        "Incluye paginas de confianza, metadatos SEO, sitemap, robots.txt, ads.txt y AdSense solo en dos paginas editoriales.",
    ])
    doc.add_heading("Limites de esta documentacion", level=2)
    doc.add_paragraph(
        "El analisis describe el codigo local y su comportamiento validado mediante las pruebas del repositorio. No equivale a una auditoria clinica, legal, de seguridad ofensiva ni a una inspeccion visual completa en navegadores reales."
    )

    doc.add_page_break()
    doc.add_heading("2. Arquitectura funcional", level=1)
    doc.add_paragraph("El sitio puede entenderse como cinco capas encadenadas:", style="Lead")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture_shape = p.add_run().add_picture(str(architecture_img), width=Inches(6.35))
    architecture_shape._inline.docPr.set("descr", "Diagrama de cinco capas: experiencia publica, acceso, cuestionarios, servicios compartidos y publicacion.")
    architecture_shape._inline.docPr.set("title", "Arquitectura funcional del sitio")
    cap = doc.add_paragraph("Figura 1. Capas funcionales observadas en el proyecto.", style="Table Citation")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Responsabilidad de cada capa", level=2)
    add_table(doc, ["Capa", "Responsabilidad principal", "Elementos clave"], [
        ("Experiencia publica", "Presentar informacion, guias, confianza y llamadas a evaluacion.", "HTML, imagenes, CSS y contenido editorial"),
        ("Acceso", "Validar audiencia, edad y confirmacion del tutor antes de enviar al test.", "formulario-de-acceso-a-tests.html + JS"),
        ("Cuestionarios", "Mostrar preguntas, capturar respuestas, calcular resultados y permitir revision.", "HTML de examenes + motores adulto/infantil"),
        ("Servicios compartidos", "Cerrar menus, manejar foco/dialogos y guardar progreso de forma tolerante a fallos.", "navigation.js, test-accessibility.js, test-storage.js"),
        ("Publicacion y control", "Exponer dominio, rastreo, monetizacion y asegurar integridad estructural.", "CNAME, sitemap, robots, ads.txt y tests"),
    ], [1500, 4100, 3760])

    doc.add_page_break()
    doc.add_heading("3. Árbol estructural del repositorio", level=1)
    doc.add_paragraph("La carpeta docs/ es el contenido publicable. El arbol siguiente muestra todos los archivos actuales, agrupados por responsabilidad.")
    tree_lines = [
        "test-neurodiverentes/",
        "+-- package.json                         # comando npm test",
        "+-- tests/                               # validacion automatizada",
        "|   +-- site-structure.test.mjs           # SEO, enlaces, recursos, a11y, HTTP",
        "|   +-- quiz-flow.test.mjs                # flujos adulto e infantil",
        "|   +-- helpers/quiz-runtime.mjs          # DOM simulado y localStorage de prueba",
        "+-- docs/                                # raiz publicada del sitio",
        "    +-- index.html                       # portada y guia introductoria",
        "    +-- test-psicologicos-gratis.html    # directorio de tests y guias",
        "    +-- formulario-de-acceso-a-tests.html # puerta previa a los tests",
        "    +-- quienes-somos.html               # identidad y autoria",
        "    +-- metodologia-de-los-cuestionarios.html # calculo y limites",
        "    +-- contacto.html                    # contacto por correo",
        "    +-- politica-de-privacidad.html      # tratamiento y privacidad",
        "    +-- CNAME                            # testneurodivergentes.com",
        "    +-- sitemap.xml                      # 12 URL indexables",
        "    +-- robots.txt                       # permite rastreo y declara sitemap",
        "    +-- ads.txt                          # autoriza el editor de AdSense",
        "    +-- assets/                          # 5 imagenes generales WebP",
        "    |   +-- tests-neurodivergentes-gratis-familia-portada-movil.webp",
        "    |   +-- tests-neurodivergentes-gratis-familia-portada-escritorio.webp",
        "    |   +-- test-neurodivergentes-480.webp",
        "    |   +-- test-neurodivergentes-768.webp",
        "    |   +-- test-neurodivergentes-1181.webp",
        "    +-- styles/                          # sistema visual",
        "    |   +-- sistema-visual.css            # base compartida y responsive",
        "    |   +-- guias-tdah.css                # articulos/guia TDAH",
        "    |   +-- tailwind-tests.min.css        # utilidades de cuestionarios",
        "    +-- JS/                              # comportamiento general",
        "    |   +-- index.js                      # acordeones y FAQ de portada",
        "    |   +-- test-psicologicos-gratis.js   # interaccion del directorio",
        "    |   +-- formulario-de-acceso-a-tests.js # validacion y enrutamiento",
        "    |   +-- TDAH.js                       # tabla/comparador de la guia",
        "    |   +-- shared/                       # servicios reutilizados",
        "    |       +-- navigation.js             # menus desktop/movil",
        "    |       +-- test-accessibility.js      # foco, Escape, Tab y movimiento",
        "    |       +-- test-storage.js            # progreso en localStorage",
        "    +-- TDA/                             # contenido y pruebas TDAH",
        "        +-- que-es-el-tdah.html            # guia principal con AdSense",
        "        +-- senales-de-tdah-en-adultos.html",
        "        +-- senales-de-tdah-en-ninos-y-cuando-consultar.html",
        "        +-- como-interpretar-un-resultado-orientativo-de-tdah.html",
        "        +-- TDAH.html                      # alias noindex con redireccion",
        "        +-- test-TDA-y-TDAH-gratuito.html  # alias noindex con redireccion",
        "        +-- assets/                        # 3 imagenes hero TDAH WebP",
        "        |   +-- hero-tdah-adulto-700.webp",
        "        |   +-- hero-tdah-adulto-1200.webp",
        "        |   +-- hero-tdah-adulto-1737.webp",
        "        +-- Examenes/                      # interfaces y motores de test",
        "            +-- test-gratuito-de-tdah-para-adultos.html",
        "            +-- adultos-tda-tdah.js",
        "            +-- test-gratuito-de-tdah-en-niños.html",
        "            +-- child-tda-tdah.js",
    ]
    add_tree(doc, tree_lines)
    doc.add_paragraph("Nota: los dos alias historicos usan noindex,follow, canonical a que-es-el-tdah.html y redireccion inmediata mediante meta refresh.", style="Table Citation")

    doc.add_heading("4 - Árbol de navegación y contenidos", level=1)
    nav_lines = [
        "Visitante",
        "+-- Portada /",
        "|   +-- Explicacion general y señales",
        "|   +-- Guias TDAH",
        "|   +-- Como funciona / metodologia / limites",
        "|   +-- Proximos pasos y preguntas frecuentes",
        "|   +-- Llamada a evaluaciones",
        "+-- Directorio /test-psicologicos-gratis.html",
        "|   +-- Test adulto -> formulario?test=TDAH&tipo=adulto",
        "|   +-- Test infantil -> formulario?test=TDAH&tipo=infantil",
        "|   +-- Guia de señales en adultos",
        "|   +-- Guia de señales en niños",
        "|   +-- Guia para interpretar resultados",
        "+-- Guia central /TDA/que-es-el-tdah.html",
        "|   +-- Informacion sobre TDA/TDAH",
        "|   +-- Comparacion interactiva",
        "|   +-- Enlaces a guias practicas",
        "|   +-- Acceso a test adulto/infantil",
        "+-- Paginas de confianza",
        "|   +-- Quienes somos",
        "|   +-- Metodologia de cuestionarios",
        "|   +-- Contacto",
        "|   +-- Politica de privacidad",
        "+-- Rutas historicas",
        "    +-- /TDA/TDAH.html -> guia central",
        "    +-- /TDA/test-TDA-y-TDAH-gratuito.html -> guia central",
    ]
    add_tree(doc, nav_lines)
    doc.add_heading("Tipos de pagina", level=2)
    add_table(doc, ["Tipo", "Cantidad", "Funcion"], [
        ("Indexables en sitemap", "12", "Contenido publico: portada, directorio, confianza, guias y dos tests."),
        ("Formulario noindex", "1", "Puerta funcional; no se promueve como contenido de busqueda."),
        ("Alias noindex", "2", "Conservan rutas antiguas y transfieren al canonical."),
        ("Con AdSense", "2", "Solo index.html y TDA/que-es-el-tdah.html."),
    ], [2300, 1100, 5960])

    doc.add_page_break()
    doc.add_heading("5. Flujo completo de una evaluacion", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    process_shape = p.add_run().add_picture(str(process_img), width=Inches(6.45))
    process_shape._inline.docPr.set("descr", "Flujo desde la entrada y seleccion del test hasta el formulario, cuestionario, calculo, resultado y acciones posteriores.")
    process_shape._inline.docPr.set("title", "Flujo principal de una evaluacion")
    cap = doc.add_paragraph("Figura 2. Proceso principal desde el descubrimiento hasta las acciones posteriores.", style="Table Citation")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Secuencia operativa", level=2)
    steps = [
        "El visitante llega desde la portada, el directorio o una guia.",
        "Selecciona el cuestionario para adultos o el cuestionario infantil.",
        "El enlace abre el formulario con parametros test=TDAH y tipo=adulto/infantil.",
        "El formulario valida nombre, genero, fecha real, edad entre 0 y 120 y audiencia. Para menores exige confirmacion del tutor.",
        "Si la seleccion y la edad son coherentes, el navegador redirige al HTML del test correspondiente.",
        "Los servicios compartidos preparan navegacion, accesibilidad y almacenamiento; despues carga el motor especifico.",
        "Cada respuesta actualiza el estado. Avanzar, retroceder, pausar o salir guarda el progreso localmente.",
        "Al completar todas las preguntas se calcula el puntaje global, las areas y el perfil orientativo.",
        "El visitante puede revisar respuestas, repetir el test o volver al directorio. El contenido recomienda valoracion profesional cuando corresponde.",
    ]
    add_numbered_list(doc, steps)

    doc.add_heading("6. Formulario previo y enrutamiento", level=1)
    doc.add_paragraph(
        "El formulario funciona como un selector controlado, no como un sistema de registro. Lee parametros de la URL, limpia el nombre del test y rechaza rutas externas con esquema o protocolo relativo."
    )
    add_table(doc, ["Entrada o regla", "Comportamiento"], [
        ("tipo=adulto", "Solo permite continuar si la edad calculada es de 18 años o mas."),
        ("tipo=infantil", "Solo permite continuar si la persona es menor de 18 y el tutor confirma."),
        ("Fecha", "Comprueba que dia, mes y año formen una fecha real y no futura."),
        ("Nombre", "Exige al menos dos caracteres; no se envia ni se guarda por este script."),
        ("Rutas", "Usa destinos internos predeterminados y descarta valores con esquema externo."),
        ("Error", "Deshabilita el boton o muestra un mensaje contextual sin enviar el formulario."),
    ], [2600, 6760])
    add_callout(doc, "Dato importante", "El formulario no contiene action hacia un servidor y su JavaScript solo utiliza los datos para validar y elegir el destino. El progreso del test se guarda despues, por separado, en localStorage.", fill="FFF7E6", accent=AMBER)

    doc.add_heading("7. Motores de cuestionarios", level=1)
    doc.add_heading("Estructura comun", level=2)
    add_bullets(doc, [
        "Inicializan la navegacion compartida antes de ejecutar el test.",
        "Convierten grupos de preguntas en una lista lineal con metadatos de seccion.",
        "Usan cinco respuestas: Nunca=0, Rara vez=1, Algunas veces=2, Frecuentemente=3 y Casi siempre=4.",
        "Mantienen currentIndex, answers, startedAt y completed en memoria y en almacenamiento local.",
        "Gestionan progreso accesible, teclado, foco, movimiento reducido, pausa, salida, revision y reinicio.",
        "Siguen funcionando aunque localStorage este bloqueado; en ese caso se pierde la recuperacion, no el test actual.",
    ])
    adult_heading = doc.add_heading("Cuestionario para adultos", level=2)
    adult_heading.paragraph_format.page_break_before = True
    add_table(doc, ["Seccion", "Preguntas", "Area"], [
        ("A", "20", "Deficit de atencion"),
        ("B", "20", "Hiperactividad e impulsividad"),
        ("C", "20", "Funciones ejecutivas"),
        ("D", "20", "Sintomas desde la infancia"),
        ("E", "20", "Entornos, deterioro y descarte"),
    ], [1100, 1500, 6760])
    doc.add_paragraph("Total: 100 preguntas; puntaje bruto maximo: 400. El perfil compara atencion, hiperactividad, infancia y presencia en varios entornos.")
    doc.add_heading("Cuestionario infantil", level=2)
    add_table(doc, ["Seccion", "Preguntas", "Area"], [
        ("A", "12", "Inatencion"),
        ("B", "12", "Hiperactividad e impulsividad"),
        ("C", "12", "Escuela y organizacion"),
        ("D", "12", "Autorregulacion y convivencia"),
        ("E", "12", "Persistencia clinica"),
    ], [1100, 1500, 6760])
    doc.add_paragraph("Total: 60 preguntas; puntaje bruto maximo: 240. Incluye aviso inicial para responsables y bloquea el boton de resultados hasta completar las 60 respuestas.")

    doc.add_heading("8. Calculo y presentacion de resultados", level=1)
    doc.add_heading("Formula", level=2)
    add_callout(doc, "Puntaje normalizado", "redondear((suma de respuestas / puntaje bruto maximo) x 100). La misma formula se aplica a cada seccion para producir valores comparables de 0 a 100.")
    bands_heading = doc.add_heading("Bandas internas", level=2)
    bands_heading.paragraph_format.page_break_before = True
    add_table(doc, ["Rango", "Etiqueta", "Color funcional"], [
        ("0-35", "Indicadores bajos", "Verde"),
        ("36-45", "Indicadores levemente elevados", "Ambar"),
        ("46-60", "Indicadores elevados relacionados con TDAH", "Naranja"),
        ("61-100", "Indicadores muy elevados relacionados con TDAH", "Rojo"),
    ], [1500, 5860, 2000])
    doc.add_heading("Interpretacion por perfil", level=2)
    doc.add_paragraph(
        "Ademas de la banda global, cada motor compara las areas principales. El adulto distingue predominio relativo inatento, hiperactivo/impulsivo o mixto y contextualiza con infancia y entornos. El infantil incorpora escuela, autorregulacion y persistencia en varios contextos."
    )
    add_callout(doc, "Limite clinico", "Las bandas son reglas internas del sitio y no puntos de corte clinicos validados. El resultado se presenta como cribado orientativo y no como diagnostico medico o psicologico.", fill="FDECEC", accent=RED)
    doc.add_heading("Salidas visibles", level=2)
    add_bullets(doc, [
        "Puntaje final sobre 100 y etiqueta de banda.",
        "Perfil orientativo redactado segun la distribucion de areas.",
        "Tarjetas con puntajes por seccion.",
        "Escala de colores y explicacion de uso responsable.",
        "Fuentes CDC/NIMH y acciones para revisar, repetir o salir.",
    ])

    doc.add_heading("9. Persistencia, privacidad y datos", level=1)
    doc.add_paragraph("No existe una base de datos del proyecto. El unico estado persistente identificado en los cuestionarios se almacena en el navegador del visitante.")
    add_table(doc, ["Clave", "Contenido", "Ciclo de vida"], [
        ("neurodivergentes:tdah-adultos:progress", "version, testId, pregunta, 100 respuestas, fechas y completed", "Se actualiza al navegar/pausar; se elimina al reiniciar"),
        ("neurodivergentes:tdah-infantil:progress", "version, testId, pregunta, 60 respuestas, fechas y completed", "Se actualiza al navegar/pausar; se elimina al reiniciar"),
    ], [3000, 3860, 2500])
    add_bullets(doc, [
        "La version actual del formato es 1; un progreso con version incompatible se descarta.",
        "Se validan longitud, valores permitidos, indice de pregunta, fechas y coherencia del estado completed.",
        "Los errores de lectura, JSON invalido o bloqueo de almacenamiento limpian/ignoran el estado sin romper la evaluacion.",
        "El formulario previo no transmite nombre, genero ni fecha a los cuestionarios y no los guarda en localStorage.",
    ])

    doc.add_heading("10. Navegacion, accesibilidad y responsive", level=1)
    doc.add_heading("Navegacion compartida", level=2)
    add_bullets(doc, [
        "navigation.js se inicializa una sola vez mediante data-site-navigation-ready.",
        "En escritorio cierra el submenu por clic externo, salida del puntero, perdida de foco o seleccion de enlace.",
        "En movil sincroniza el elemento details, bloquea visualmente el fondo mediante la clase menu-open y permite cerrar con Escape.",
        "Al cambiar a un ancho de escritorio, fuerza el cierre del menu movil para evitar estados inconsistentes.",
    ])
    doc.add_heading("Accesibilidad de cuestionarios", level=2)
    add_bullets(doc, [
        "Barra de progreso con role=progressbar, aria-valuenow y aria-valuetext.",
        "Estado de preguntas y resultados con regiones aria-live.",
        "Opciones operables como radios y navegables por flechas, Inicio y Fin.",
        "Dialogos de pausa y aviso infantil con foco controlado, Tab ciclico y Escape.",
        "Respeto a prefers-reduced-motion para sustituir desplazamiento suave por inmediato.",
        "Menus especificos para escritorio y movil; CSS con reglas de altura segura usando 100svh.",
    ])
    styles_heading = doc.add_heading("Estilos", level=2)
    styles_heading.paragraph_format.page_break_before = True
    add_table(doc, ["Archivo", "Uso"], [
        ("sistema-visual.css", "Sistema comun, footer responsable, menus, componentes y responsive."),
        ("guias-tdah.css", "Presentacion editorial de las tres guias practicas."),
        ("tailwind-tests.min.css", "Clases utilitarias compiladas usadas por las interfaces de cuestionario."),
    ], [2800, 6560])

    doc.add_heading("11. SEO, publicacion y monetizacion", level=1)
    doc.add_heading("Publicacion", level=2)
    add_bullets(doc, [
        "docs/CNAME fija el dominio testneurodivergentes.com.",
        "robots.txt permite el rastreo general y declara la ubicacion del sitemap.",
        "sitemap.xml contiene exactamente las 12 paginas indexables esperadas por las pruebas.",
        "Cada pagina indexable tiene titulo, descripcion, canonical absoluto y un unico H1.",
        "Las rutas antiguas son noindex,follow y apuntan al canonical actual para evitar competir como duplicados.",
    ])
    doc.add_heading("AdSense", level=2)
    add_bullets(doc, [
        "El archivo ads.txt autoriza pub-5855172805291728 como vendedor directo de Google.",
        "El cargador de AdSense aparece unicamente en index.html y TDA/que-es-el-tdah.html.",
        "Los tests, formulario, paginas legales, directorio y guias secundarias no cargan el script publicitario.",
        "La captura administrativa revisada por el usuario mostraba ads.txt como Autorizado y el sitio en Preparando; ese estado externo no forma parte del codigo del repositorio.",
    ])
    add_callout(doc, "Separacion de responsabilidades", "El repositorio controla que el codigo y ads.txt esten publicados. La aprobacion final y el servicio de anuncios dependen de la revision externa de Google AdSense.", fill="FFF7E6", accent=AMBER)

    doc.add_heading("12. Pruebas y control de calidad", level=1)
    doc.add_paragraph("El comando npm test ejecuta Node Test Runner de forma secuencial sobre dos grupos de pruebas.")
    add_table(doc, ["Grupo", "Cobertura"], [
        ("Flujo de cuestionarios", "Minimo/maximo, limites de bandas, guardado, recarga, version incompatible, reinicio, pausa, responsive, teclado y bloqueo del resultado infantil."),
        ("Estructura del sitio", "Metadatos, sitemap, imagenes, IDs, enlaces externos, footer, autoria, metodologia, guias, AdSense, recursos, orden de scripts, navegacion, accesibilidad, responsive y HTTP sin 404."),
    ], [2600, 6760])
    add_callout(doc, "Estado actual", "36 pruebas ejecutadas, 36 aprobadas, 0 fallos, 0 omitidas. Verificacion realizada el 10 de agosto de 2026.", fill="EDF7F2", accent=GREEN)
    doc.add_heading("Que garantizan", level=2)
    add_bullets(doc, [
        "Que los flujos principales responden como espera la logica implementada.",
        "Que todos los enlaces y recursos locales existen y responden por un servidor HTTP de prueba.",
        "Que el sitemap coincide con los canonicals indexables.",
        "Que no se duplican scripts y las dependencias compartidas cargan antes del motor del test.",
    ])
    doc.add_heading("Que no garantizan por si solas", level=2)
    add_bullets(doc, [
        "Aprobacion de AdSense, posicionamiento SEO o disponibilidad del proveedor de hosting.",
        "Exactitud clinica de las preguntas y bandas internas.",
        "Compatibilidad visual perfecta en todos los dispositivos y navegadores reales.",
        "Proteccion frente a todos los ataques posibles o cumplimiento legal integral.",
    ])

    doc.add_heading("13. Mantenimiento y puntos de atencion", level=1)
    doc.add_heading("Procedimiento recomendado para cambios", level=2)
    add_numbered_list(doc, [
        "Modificar el archivo fuente dentro de docs/ y conservar rutas relativas coherentes.",
        "Si se agrega una pagina indexable, añadir canonical, descripcion, H1 unico y entrada en sitemap.xml.",
        "Si se agrega JavaScript a un test, cargar navigation, accessibility y storage antes del motor especifico.",
        "Si cambia la estructura de progreso, aumentar la version y adaptar las pruebas de compatibilidad.",
        "Ejecutar npm test y git diff --check antes de publicar.",
        "Verificar en navegador real los cambios visuales o de interaccion de alto impacto.",
    ])
    doc.add_heading("Puntos de atencion observados", level=2)
    add_bullets(doc, [
        "El sitio es estatico: no existe panel propio, API ni base de datos para centralizar resultados.",
        "El contenido de varias paginas incluye CSS extenso dentro del HTML; cambios visuales globales pueden requerir auditoria pagina por pagina.",
        "El formulario previo solicita datos personales para enrutamiento, aunque el script actual no los guarda. Cualquier futura integracion de analitica o backend debe reevaluar privacidad y consentimiento.",
        "La funcion infantil calcula una propiedad interna de subtipo, pero la plantilla visible actual presenta perfil y areas; cualquier cambio de etiqueta diagnostica debe revisarse con especial cautela.",
        "Las bandas y cuestionarios se documentan como orientativos y sin validacion clinica publicada; esa advertencia debe preservarse.",
        "Las dos rutas historicas deben mantenerse mientras existan enlaces externos que puedan depender de ellas.",
    ])

    doc.add_page_break()
    doc.add_heading("Anexo A. Inventario funcional de archivos", level=1)
    inventory = [
        ("docs/index.html", "Portada, contenido pilar, FAQ, enlaces a guias/tests y AdSense."),
        ("docs/test-psicologicos-gratis.html", "Directorio de los dos tests y tres guias practicas."),
        ("docs/formulario-de-acceso-a-tests.html", "Formulario noindex previo a la evaluacion."),
        ("docs/quienes-somos.html", "Identidad, autoria y transparencia."),
        ("docs/metodologia-de-los-cuestionarios.html", "Formula, rangos y limites metodologicos."),
        ("docs/contacto.html", "Canal mailto de contacto."),
        ("docs/politica-de-privacidad.html", "Informacion de privacidad y contacto."),
        ("docs/TDA/que-es-el-tdah.html", "Guia central, comparador, enlaces y AdSense."),
        ("docs/TDA/senales-de-tdah-en-adultos.html", "Guia practica para adultos."),
        ("docs/TDA/senales-de-tdah-en-ninos-y-cuando-consultar.html", "Guia practica infantil."),
        ("docs/TDA/como-interpretar-un-resultado-orientativo-de-tdah.html", "Guia posterior al test."),
        ("docs/TDA/TDAH.html", "Alias noindex y redireccion a la guia central."),
        ("docs/TDA/test-TDA-y-TDAH-gratuito.html", "Segundo alias noindex y redireccion."),
        ("docs/TDA/Examenes/test-gratuito-de-tdah-para-adultos.html", "Interfaz accesible del test adulto."),
        ("docs/TDA/Examenes/adultos-tda-tdah.js", "100 preguntas, estado, puntuacion y resultados adultos."),
        ("docs/TDA/Examenes/test-gratuito-de-tdah-en-niños.html", "Interfaz accesible, aviso del tutor y resultado bloqueado."),
        ("docs/TDA/Examenes/child-tda-tdah.js", "60 preguntas, estado, puntuacion y resultados infantiles."),
        ("docs/JS/formulario-de-acceso-a-tests.js", "Edad, consentimiento, parametros y redireccion segura."),
        ("docs/JS/index.js", "Acordeones moviles y expansion de FAQ."),
        ("docs/JS/test-psicologicos-gratis.js", "Interacciones del directorio."),
        ("docs/JS/TDAH.js", "Comparador interactivo de la guia central."),
        ("docs/JS/shared/navigation.js", "Menus desktop/movil e inicializacion idempotente."),
        ("docs/JS/shared/test-accessibility.js", "Foco, teclado, dialogs y movimiento reducido."),
        ("docs/JS/shared/test-storage.js", "Persistencia validada y tolerante a fallos."),
        ("docs/styles/sistema-visual.css", "Estilo comun y responsive."),
        ("docs/styles/guias-tdah.css", "Estilo editorial de guias."),
        ("docs/styles/tailwind-tests.min.css", "Utilidades compiladas de los tests."),
        ("docs/assets/*.webp", "Cinco variantes responsivas generales."),
        ("docs/TDA/assets/*.webp", "Tres variantes responsivas del hero TDAH."),
        ("docs/CNAME", "Dominio personalizado."),
        ("docs/sitemap.xml", "Inventario de URL indexables."),
        ("docs/robots.txt", "Politica de rastreo."),
        ("docs/ads.txt", "Autorizacion del editor publicitario."),
        ("tests/site-structure.test.mjs", "Pruebas estructurales, editoriales, SEO, a11y y HTTP."),
        ("tests/quiz-flow.test.mjs", "Pruebas de flujo de ambos cuestionarios."),
        ("tests/helpers/quiz-runtime.mjs", "Entorno simulado de navegador para tests."),
        ("package.json", "Define npm test y el orden de ejecucion."),
    ]
    add_table(doc, ["Archivo o grupo", "Responsabilidad"], inventory, [3500, 5860])
    doc.add_paragraph("Fuente de esta documentacion: inspeccion del repositorio local y ejecucion de npm test el 10 de agosto de 2026.", style="Table Citation")

    # Core metadata.
    props = doc.core_properties
    props.title = "Arbol de procesos y estructura de testneurodivergentes.com"
    props.subject = "Documentacion tecnica y funcional del proyecto"
    props.author = "OpenAI Codex"
    props.keywords = "TDAH, arquitectura, procesos, estructura, sitio estatico, cuestionarios"
    props.comments = "Generado a partir de una auditoria del repositorio local."

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_document())
